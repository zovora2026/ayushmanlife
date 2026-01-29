"""
AyushmanLife SmartClaims™ - Core Claim Processing Logic

AI-powered insurance claims automation for Indian hospitals.
Handles document extraction, AI analysis, ICD-10/CPT coding,
completeness validation, and discharge summary generation.
"""

from __future__ import annotations

import io
import json
import re
from dataclasses import asdict, dataclass, field
from typing import Any

from PyPDF2 import PdfReader
from dotenv import load_dotenv
from openai import OpenAI
from PIL import Image
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
from docx import Document
import pytesseract

load_dotenv()


# ---------------------------------------------------------------------------
# Data structures
# ---------------------------------------------------------------------------


@dataclass
class ClaimData:
    """Structured claim information."""

    patient_name: str = ""
    patient_age: str = ""
    patient_gender: str = ""
    hospital_name: str = ""
    treatment_procedure: str = ""
    claim_amount: str = ""
    admission_date: str = ""
    discharge_date: str = ""
    insurance_scheme: str = ""
    clinical_notes: str = ""
    extracted_text: str = ""


@dataclass
class CompletenessItem:
    """Single completeness checklist item."""

    item: str
    present: bool
    notes: str = ""


@dataclass
class AnalysisResult:
    """AI analysis output."""

    patient_summary: str = ""
    icd10_codes: list[str] = field(default_factory=list)
    cpt_codes: list[str] = field(default_factory=list)
    completeness: list[CompletenessItem] = field(default_factory=list)
    missing_documents: list[str] = field(default_factory=list)
    rejection_risk: float = 0.0
    recommendations: list[str] = field(default_factory=list)
    approval_likelihood: str = ""
    raw_ai_response: str = ""


# ---------------------------------------------------------------------------
# Document extraction
# ---------------------------------------------------------------------------


def extract_text_from_pdf(file_content: bytes, filename: str = "document.pdf") -> str:
    """
    Extract text from a PDF file using PyPDF2.

    Args:
        file_content: Raw PDF bytes.
        filename: Optional filename for error messages.

    Returns:
        Extracted text. Empty string if extraction fails.
    """
    try:
        reader = PdfReader(io.BytesIO(file_content))
        chunks = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                chunks.append(text.strip())
        return "\n\n".join(chunks) if chunks else ""
    except Exception as e:
        raise ValueError(f"PDF extraction failed for {filename}: {e}") from e


def extract_text_from_image(
    file_content: bytes, filename: str = "image.jpg", lang: str = "eng"
) -> str:
    """
    Extract text from an image using OCR (Tesseract).

    Suitable for handwritten prescriptions and scanned documents.

    Args:
        file_content: Raw image bytes.
        filename: Optional filename for error messages.
        lang: Tesseract language code (e.g. 'eng', 'hin', 'hin+eng').

    Returns:
        Extracted text. Empty string if OCR fails.
    """
    try:
        image = Image.open(io.BytesIO(file_content))
        if image.mode not in ("L", "RGB"):
            image = image.convert("RGB")
        try:
            return pytesseract.image_to_string(image, lang=lang).strip() or ""
        except pytesseract.TesseractError:
            if lang != "eng":
                return pytesseract.image_to_string(image, lang="eng").strip() or ""
            raise
    except Exception as e:
        raise ValueError(f"OCR extraction failed for {filename}: {e}") from e


def transcribe_audio(audio_bytes: bytes, *, api_key: str, filename: str = "audio.webm") -> str:
    """
    Transcribe audio to text using OpenAI Whisper.

    Args:
        audio_bytes: Raw audio file bytes (e.g. webm, mp3, wav).
        api_key: OpenAI API key.
        filename: Hint for format (extension used).

    Returns:
        Transcribed text.
    """
    client = OpenAI(api_key=api_key)
    f = io.BytesIO(audio_bytes)
    f.name = filename
    r = client.audio.transcriptions.create(model="whisper-1", file=f)
    return (r.text or "").strip()


def extract_text_from_file(
    file_content: bytes, filename: str, ocr_lang: str = "eng"
) -> str:
    """
    Route to appropriate extractor based on file type.

    Supports: PDF, JPG/JPEG, PNG, TXT.
    """
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        return extract_text_from_pdf(file_content, filename)
    if name.endswith((".jpg", ".jpeg", ".png")):
        return extract_text_from_image(file_content, filename, ocr_lang)
    if name.endswith(".txt"):
        return file_content.decode("utf-8", errors="replace").strip()
    raise ValueError(f"Unsupported file type: {filename}")


# ---------------------------------------------------------------------------
# AI analysis (OpenAI)
# ---------------------------------------------------------------------------


def _build_claim_prompt(claim: ClaimData, language: str = "English") -> str:
    """Build the analysis prompt from claim data."""
    lang_note = "Respond in Hindi." if language.lower() == "hindi" else "Respond in English."
    return f"""Analyze this Indian hospital insurance claim for Ayushman Bharat / CGHS / ECHS:

Patient: {claim.patient_name}, {claim.patient_age}, {claim.patient_gender}
Hospital: {claim.hospital_name}
Treatment/Procedure: {claim.treatment_procedure}
Claim Amount: ₹{claim.claim_amount}
Admission: {claim.admission_date} | Discharge: {claim.discharge_date}
Insurance Scheme: {claim.insurance_scheme}

Clinical Notes / Extracted Text:
{claim.clinical_notes or "(none)"}

{lang_note}

Provide a structured analysis with exactly these sections (use clear headers):

1. **Patient Details Summary**: 2–3 sentence summary of the case.

2. **ICD-10 Codes**: List appropriate ICD-10 codes (one per line), briefly justified.

3. **CPT / Procedure Codes**: List relevant procedure codes with short descriptions.

4. **Completeness Checklist** (10 items): For each item, state YES or NO and a brief note.
   Items: (a) Discharge summary, (b) Final diagnosis, (c) Operation notes (if surgery),
   (d) Investigation reports, (e) Pharmacy bills, (f) Implant/stent details (if any),
   (g) Pre-authorization approval, (h) ID proof, (i) Insurance card/copy, (j) Consent forms.

5. **Missing Documents**: List any missing documents or information.

6. **Rejection Risk**: A single percentage (0–100) and one sentence explanation.

7. **Recommendations**: Specific fixes or improvements needed (bullet points).

8. **Approval Likelihood**: Low / Medium / High with one sentence reasoning."""


def analyze_claim_with_ai(
    claim: ClaimData,
    *,
    api_key: str,
    model: str = "gpt-4",
    language: str = "English",
) -> AnalysisResult:
    """
    Run full AI analysis on the claim.

    Uses OpenAI API. Requires valid api_key.
    """
    client = OpenAI(api_key=api_key)
    prompt = _build_claim_prompt(claim, language)

    response = client.chat.completions.create(
        model=model,
        messages=[
            {
                "role": "system",
                "content": "You are an expert Indian health insurance claims analyst. "
                "You understand Ayushman Bharat, CGHS, ECHS, and private schemes. "
                "Give accurate, actionable analysis.",
            },
            {"role": "user", "content": prompt},
        ],
        temperature=0.2,
    )
    raw = response.choices[0].message.content or ""

    return _parse_ai_response(raw, claim)


def _parse_ai_response(raw: str, claim: ClaimData) -> AnalysisResult:
    """Parse raw AI text into structured AnalysisResult."""
    result = AnalysisResult(raw_ai_response=raw)

    # Patient summary
    m = re.search(
        r"\*\*Patient Details Summary\*\*[\s:]*\n(.+?)(?=\n\s*\*\*|\n\n\n|\Z)",
        raw,
        re.DOTALL | re.IGNORECASE,
    )
    if m:
        result.patient_summary = m.group(1).strip()

    # ICD-10
    m = re.search(
        r"\*\*ICD-10 Codes\*\*[\s:]*\n(.+?)(?=\n\s*\*\*|\n\n\n|\Z)",
        raw,
        re.DOTALL | re.IGNORECASE,
    )
    if m:
        block = m.group(1)
        for line in block.split("\n"):
            line = line.strip()
            if not line or line.lower().startswith("no "):
                continue
            code = re.search(r"[A-Z]\d{2}(?:\.\d{2,4})?", line, re.IGNORECASE)
            if code:
                result.icd10_codes.append(line)

    # CPT
    m = re.search(
        r"\*\*CPT\s*/\s*Procedure Codes\*\*[\s:]*\n(.+?)(?=\n\s*\*\*|\n\n\n|\Z)",
        raw,
        re.DOTALL | re.IGNORECASE,
    )
    if m:
        block = m.group(1)
        for line in block.split("\n"):
            line = line.strip()
            if line and not line.lower().startswith("no "):
                result.cpt_codes.append(line)

    # Missing documents
    m = re.search(
        r"\*\*Missing Documents\*\*[\s:]*\n(.+?)(?=\n\s*\*\*|\n\n\n|\Z)",
        raw,
        re.DOTALL | re.IGNORECASE,
    )
    if m:
        block = m.group(1)
        for line in block.split("\n"):
            line = line.strip().strip("-*•")
            if line and line.lower() not in ("none", "n/a", "-"):
                result.missing_documents.append(line)

    # Rejection risk
    m = re.search(r"\*\*Rejection Risk\*\*[\s:]*\n?.+?(\d{1,3})\s*%", raw, re.IGNORECASE)
    if m:
        result.rejection_risk = min(100.0, max(0.0, float(m.group(1))))

    # Recommendations
    m = re.search(
        r"\*\*Recommendations\*\*[\s:]*\n(.+?)(?=\n\s*\*\*|\n\n\n|\Z)",
        raw,
        re.DOTALL | re.IGNORECASE,
    )
    if m:
        block = m.group(1)
        for line in block.split("\n"):
            line = line.strip().strip("-*•")
            if line:
                result.recommendations.append(line)

    # Approval likelihood
    for phrase in ["approval likelihood", "approval likelihood:"]:
        idx = raw.lower().find(phrase)
        if idx >= 0:
            rest = raw[idx + len(phrase) : idx + len(phrase) + 200]
            lik = re.search(r"\b(Low|Medium|High)\b", rest, re.IGNORECASE)
            if lik:
                result.approval_likelihood = lik.group(1)
                break

    # Default 10-point completeness from raw or template
    default_items = [
        "Discharge summary",
        "Final diagnosis",
        "Operation notes (if surgery)",
        "Investigation reports",
        "Pharmacy bills",
        "Implant/stent details (if any)",
        "Pre-authorization approval",
        "ID proof",
        "Insurance card/copy",
        "Consent forms",
    ]
    result.completeness = [
        CompletenessItem(item=item, present=False, notes="")
        for item in default_items
    ]
    comp_block = re.search(
        r"\*\*Completeness Checklist\*\*[\s:]*\n(.+?)(?=\n\s*\*\*|\n\n\n|\Z)",
        raw,
        re.DOTALL | re.IGNORECASE,
    )
    if comp_block:
        text = comp_block.group(1).lower()
        for i, item in enumerate(default_items):
            if i < len(result.completeness):
                kw = item.split("(")[0].strip().lower()
                result.completeness[i].present = "yes" in text and kw in text

    return result


def generate_icd10_codes(
    clinical_text: str,
    *,
    api_key: str,
    model: str = "gpt-4",
) -> list[str]:
    """Generate ICD-10 code suggestions from clinical text."""
    client = OpenAI(api_key=api_key)
    r = client.chat.completions.create(
        model=model,
        messages=[
            {
                "role": "system",
                "content": "You are an expert medical coder. Reply with only ICD-10 codes, one per line.",
            },
            {"role": "user", "content": f"Suggest ICD-10 codes for:\n\n{clinical_text}"},
        ],
        temperature=0.1,
    )
    text = r.choices[0].message.content or ""
    return [x.strip() for x in text.split("\n") if x.strip() and re.search(r"[A-Z]\d{2}", x, re.I)]


def generate_cpt_codes(
    procedure_text: str,
    *,
    api_key: str,
    model: str = "gpt-4",
) -> list[str]:
    """Generate CPT / procedure code suggestions."""
    client = OpenAI(api_key=api_key)
    r = client.chat.completions.create(
        model=model,
        messages=[
            {
                "role": "system",
                "content": "You are an expert medical coder. Suggest CPT/procedure codes for Indian hospital claims.",
            },
            {"role": "user", "content": f"Procedure/treatment: {procedure_text}"},
        ],
        temperature=0.1,
    )
    text = r.choices[0].message.content or ""
    return [x.strip() for x in text.split("\n") if x.strip()]


# ---------------------------------------------------------------------------
# Completeness & risk
# ---------------------------------------------------------------------------

COMPLETENESS_CHECKLIST = [
    "Discharge summary",
    "Final diagnosis",
    "Operation notes (if surgery)",
    "Investigation reports",
    "Pharmacy bills",
    "Implant/stent details (if any)",
    "Pre-authorization approval",
    "ID proof",
    "Insurance card/copy",
    "Consent forms",
]


def check_completeness(
    claim: ClaimData,
    analysis: AnalysisResult | None = None,
) -> list[CompletenessItem]:
    """
    Validate claim against 10-point completeness checklist.

    Uses analysis completeness if provided; otherwise derives from claim + defaults.
    """
    if analysis and analysis.completeness:
        return analysis.completeness

    items: list[CompletenessItem] = []
    combined = f"{claim.clinical_notes} {claim.extracted_text}".lower()

    for name in COMPLETENESS_CHECKLIST:
        kw = name.split("(")[0].strip().lower().replace(" ", " ")
        present = kw in combined or any(w in combined for w in kw.split())
        items.append(CompletenessItem(item=name, present=bool(present), notes=""))
    return items


def calculate_rejection_risk(
    claim: ClaimData,
    analysis: AnalysisResult | None = None,
    completeness: list[CompletenessItem] | None = None,
) -> float:
    """
    Compute rejection risk score (0–100).

    Uses AI-derived risk if available; otherwise heuristic from completeness.
    """
    if analysis and analysis.rejection_risk > 0:
        return analysis.rejection_risk

    comp = completeness or check_completeness(claim, analysis)
    missing = sum(1 for c in comp if not c.present)
    # ~10% per missing item, capped at 90%
    return min(90.0, missing * 10.0)


# ---------------------------------------------------------------------------
# Discharge summary (FHIR R4–style)
# ---------------------------------------------------------------------------


def generate_discharge_summary(
    claim: ClaimData,
    analysis: AnalysisResult,
    format: str = "fhir",
) -> str:
    """
    Generate a discharge summary.

    format: "fhir" -> FHIR R4–style JSON; "text" -> plain text.
    """
    if format == "text":
        return _discharge_summary_text(claim, analysis)

    # FHIR R4–style structure
    fhir = {
        "resourceType": "Composition",
        "status": "final",
        "type": {
            "coding": [{"system": "http://loinc.org", "code": "18842-5", "display": "Discharge summary"}]
        },
        "subject": {
            "display": claim.patient_name,
            "extension": [
                {"url": "age", "valueString": claim.patient_age},
                {"url": "gender", "valueString": claim.patient_gender},
            ],
        },
        "date": claim.discharge_date or "",
        "title": "Discharge Summary",
        "section": [
            {
                "title": "Patient Information",
                "text": {
                    "status": "generated",
                    "div": f"<div>Name: {claim.patient_name}; Age: {claim.patient_age}; Gender: {claim.patient_gender}; Hospital: {claim.hospital_name}</div>",
                },
            },
            {
                "title": "Clinical Summary",
                "text": {"status": "generated", "div": f"<div>{analysis.patient_summary or 'N/A'}</div>"},
            },
            {
                "title": "Treatment / Procedure",
                "text": {"status": "generated", "div": f"<div>{claim.treatment_procedure}</div>"},
            },
            {
                "title": "ICD-10 Codes",
                "text": {
                    "status": "generated",
                    "div": f"<div>{', '.join(analysis.icd10_codes) or 'N/A'}</div>",
                },
            },
            {
                "title": "CPT / Procedure Codes",
                "text": {
                    "status": "generated",
                    "div": f"<div>{', '.join(analysis.cpt_codes) or 'N/A'}</div>",
                },
            },
            {
                "title": "Admission / Discharge",
                "text": {
                    "status": "generated",
                    "div": f"<div>Admission: {claim.admission_date}; Discharge: {claim.discharge_date}</div>",
                },
            },
        ],
    }
    return json.dumps(fhir, indent=2)


def _discharge_summary_text(claim: ClaimData, analysis: AnalysisResult) -> str:
    lines = [
        "DISCHARGE SUMMARY",
        "==================",
        "",
        f"Patient: {claim.patient_name} | Age: {claim.patient_age} | Gender: {claim.patient_gender}",
        f"Hospital: {claim.hospital_name}",
        f"Admission: {claim.admission_date} | Discharge: {claim.discharge_date}",
        "",
        "Treatment / Procedure:",
        claim.treatment_procedure or "N/A",
        "",
        "Clinical Summary:",
        analysis.patient_summary or "N/A",
        "",
        "ICD-10 Codes: " + ", ".join(analysis.icd10_codes) or "N/A",
        "CPT Codes: " + ", ".join(analysis.cpt_codes) or "N/A",
    ]
    return "\n".join(lines)


# ---------------------------------------------------------------------------
# Export helpers
# ---------------------------------------------------------------------------


def export_claim_json(claim: ClaimData, analysis: AnalysisResult) -> str:
    """Structured claim + analysis as JSON."""
    out: dict[str, Any] = {
        "claim": asdict(claim),
        "analysis": {
            "patient_summary": analysis.patient_summary,
            "icd10_codes": analysis.icd10_codes,
            "cpt_codes": analysis.cpt_codes,
            "missing_documents": analysis.missing_documents,
            "rejection_risk": analysis.rejection_risk,
            "recommendations": analysis.recommendations,
            "approval_likelihood": analysis.approval_likelihood,
            "completeness": [{"item": c.item, "present": c.present, "notes": c.notes} for c in analysis.completeness],
        },
    }
    return json.dumps(out, indent=2)


def _escape_html(s: str) -> str:
    """Escape for ReportLab Paragraph (basic HTML)."""
    if not s:
        return ""
    return (
        s.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
    )


def create_analysis_pdf(
    claim: ClaimData,
    analysis: AnalysisResult,
    output_buffer: io.BytesIO,
) -> None:
    """Write analysis report PDF to output_buffer."""
    doc = SimpleDocTemplate(output_buffer, pagesize=A4, rightMargin=inch, leftMargin=inch)
    styles = getSampleStyleSheet()
    story = []

    story.append(Paragraph("AyushmanLife SmartClaims™ – Analysis Report", styles["Title"]))
    story.append(Spacer(1, 0.25 * inch))

    story.append(Paragraph("Patient &amp; Claim", styles["Heading2"]))
    story.append(Paragraph(f"<b>Patient:</b> {_escape_html(claim.patient_name)}, {_escape_html(claim.patient_age)}, {_escape_html(claim.patient_gender)}", styles["Normal"]))
    story.append(Paragraph(f"<b>Hospital:</b> {_escape_html(claim.hospital_name)}", styles["Normal"]))
    story.append(Paragraph(f"<b>Treatment:</b> {_escape_html(claim.treatment_procedure)}", styles["Normal"]))
    story.append(Paragraph(f"<b>Amount:</b> ₹{_escape_html(claim.claim_amount)}", styles["Normal"]))
    story.append(Spacer(1, 0.2 * inch))

    story.append(Paragraph("Summary", styles["Heading2"]))
    story.append(Paragraph(_escape_html(analysis.patient_summary or "N/A"), styles["Normal"]))
    story.append(Spacer(1, 0.2 * inch))

    story.append(Paragraph("ICD-10 Codes", styles["Heading2"]))
    for c in analysis.icd10_codes:
        story.append(Paragraph(f"• {_escape_html(c)}", styles["Normal"]))
    story.append(Spacer(1, 0.2 * inch))

    story.append(Paragraph("CPT / Procedure Codes", styles["Heading2"]))
    for c in analysis.cpt_codes:
        story.append(Paragraph(f"• {_escape_html(c)}", styles["Normal"]))
    story.append(Spacer(1, 0.2 * inch))

    story.append(Paragraph("Completeness", styles["Heading2"]))
    for c in analysis.completeness:
        status = "✓" if c.present else "✗"
        story.append(Paragraph(f"{status} {_escape_html(c.item)}", styles["Normal"]))
    story.append(Spacer(1, 0.2 * inch))

    story.append(Paragraph("Missing Documents", styles["Heading2"]))
    for m in analysis.missing_documents:
        story.append(Paragraph(f"• {_escape_html(m)}", styles["Normal"]))
    story.append(Spacer(1, 0.2 * inch))

    story.append(Paragraph(f"Rejection Risk: {analysis.rejection_risk:.0f}%", styles["Heading2"]))
    story.append(Paragraph(f"Approval Likelihood: {_escape_html(analysis.approval_likelihood or 'N/A')}", styles["Normal"]))
    story.append(Spacer(1, 0.2 * inch))

    story.append(Paragraph("Recommendations", styles["Heading2"]))
    for r in analysis.recommendations:
        story.append(Paragraph(f"• {_escape_html(r)}", styles["Normal"]))

    doc.build(story)


def create_discharge_summary_docx(
    claim: ClaimData,
    analysis: AnalysisResult,
    output_buffer: io.BytesIO,
) -> None:
    """Write discharge summary as DOCX to output_buffer."""
    doc = Document()
    doc.add_heading("Discharge Summary", 0)
    doc.add_paragraph(f"Patient: {claim.patient_name} | Age: {claim.patient_age} | Gender: {claim.patient_gender}")
    doc.add_paragraph(f"Hospital: {claim.hospital_name}")
    doc.add_paragraph(f"Admission: {claim.admission_date} | Discharge: {claim.discharge_date}")
    doc.add_paragraph("")
    doc.add_heading("Treatment / Procedure", level=1)
    doc.add_paragraph(claim.treatment_procedure or "N/A")
    doc.add_heading("Clinical Summary", level=1)
    doc.add_paragraph(analysis.patient_summary or "N/A")
    doc.add_heading("ICD-10 Codes", level=1)
    doc.add_paragraph(", ".join(analysis.icd10_codes) or "N/A")
    doc.add_heading("CPT Codes", level=1)
    doc.add_paragraph(", ".join(analysis.cpt_codes) or "N/A")
    doc.save(output_buffer)
